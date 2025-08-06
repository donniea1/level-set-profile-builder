from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import List
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

app = FastAPI()
SAVE_DIR = Path("generated_docs")
SAVE_DIR.mkdir(parents=True, exist_ok=True)

class ProfileSections(BaseModel):
    step1: str
    step2: str
    step3: List[List[str]]
    step4: List[List[str]]
    step5: List[str]
    step6: List[List[str]]
    summary: str

class ProfileRequest(BaseModel):
    client_name: str
    profile_sections: ProfileSections

@app.post("/generate-docx")
def generate_docx(data: ProfileRequest):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    title = doc.add_paragraph(f"Level Set Profile Report: {data.client_name}")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)

    def add_header(text):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_table(headers, rows):
        table = doc.add_table(rows=1, cols=len(headers))
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            run = hdr_cells[i].paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(11)
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), "261418")
            hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
            run.font.color.rgb = RGBColor(255, 255, 255)
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for row in rows:
            cells = table.add_row().cells
            for i, cell in enumerate(row):
                cells[i].text = str(cell)
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    add_header("🌱 Career Alignment Assessment")
    doc.add_paragraph(data.profile_sections.step1)

    add_header("🏷 Current Functional Domain")
    doc.add_paragraph(data.profile_sections.step2)

    add_header("📈 Trajectory Insight")
    add_table(["Stage", "Title", "Years"], data.profile_sections.step3)

    add_header("💰 Compensation Intelligence")
    add_table(["Stage", "Typical Salary"], data.profile_sections.step4)

    add_header("🔑 Strategic Readiness: Key Growth Levers")
    for item in data.profile_sections.step5:
        doc.add_paragraph(f"- {item}")

    add_header("🌟 Future-Facing Roles")
    add_table(["Title", "Executive?", "Salary"], data.profile_sections.step6)

    add_header("❤️ Readiness Summary")
    doc.add_paragraph(data.profile_sections.summary)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "© 2025 LEVEL SET NEXT | Pivot With Power    |    Phone: 678-870-8341"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    filename = f"{data.client_name.replace(' ', '_')}_LevelSetProfile.docx"
    file_path = SAVE_DIR / filename
    doc.save(file_path)
    return {"download_url": f"/download/{filename}"}

@app.get("/download/{filename}")
def download_file(filename: str):
    file_path = SAVE_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found.")
    return FileResponse(path=file_path, filename=filename, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
